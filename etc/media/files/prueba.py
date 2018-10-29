from docx import *
import codecs


#document = Document('libres.docx')

def contadorTXT(nombre):
	document = Document('libres.docx')
	total=0
	for p in document.paragraphs:
		total=total+len(p.text.split())
		print (str(p.text.split())+":"+str(len(p.text.split())))

	#print (total)

	
   


from django.conf import settings
import os
from docx import Document

def handle_uploaded_file(f):
    with open((os.path.join(settings.MEDIA_ROOT,"files",f.name)),'wb+') as destination:
        for chunk in f.chunks():
            destination.write(chunk)


#Contador de palabras, solo para archivos docx
def contadorPalabras(nombre):
	total=0
	try:
		document = Document((os.path.join(settings.MEDIA_ROOT,"files",nombre)))
		for p in document.paragraphs:
			total=total+len(p.text.split())
	except:
		total=contadorPalabrasAux(nombre)
	return total

#Metodo de respaldo en caso de que el contadorPalabras Falle
def contadorPalabrasAux(nombre):
	print ("Ruta")
	print((os.path.join(settings.MEDIA_ROOT,"files",nombre)))
	with open((os.path.join(settings.MEDIA_ROOT,"files",nombre)), 'rb') as f:
		source_stream = io.BytesIO(f.read())
	document = Document(source_stream)
	total=0
	for p in document.paragraphs:
		total=total+len(p.text.split())
	source_stream.close()
	return total

#Metodo contador de palabras, si el contenido a traducir es ingresado por la caja de texto
def contarTXT(string):
	char=0
	word=1
	for i in string:
		char=char+1
		if(i==' '):
			word=word+1
	
	return word
#Metodo que crea un fichero txt con el titulo de la cotizacion si el contenido a traducir proviene de la caja de texto
def crearTXT(texto,nombre):
	f = open (settings.MEDIA_ROOT+settings.FILE_URL+str(nombre),'w')
	f.write(texto)
	f.close()